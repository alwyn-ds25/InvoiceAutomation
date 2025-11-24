from typing import List, Dict, Literal, Optional, TypedDict
from pydantic import BaseModel, Field
import pdfplumber
import docx

# --- Data Models for OCR Output ---

class PageResult(TypedDict):
    page_number: int
    text: str

class TableResult(TypedDict):
    page_number: int
    cells: List[List[str]]

class OCRResult(BaseModel):
    """
    Standardized output for the cascading OCR process.
    """
    status: Literal["OCR_DONE", "FAILED_OCR"]
    avg_confidence: float = Field(..., ge=0.0, le=1.0)
    pages: List[PageResult]
    tables: List[TableResult]
    raw_engine_trace: Dict


# --- Mock/Placeholder OCR Engine Functions ---
# In a real implementation, these would call external APIs or libraries.

def try_typhoon_ocr(image_paths: list[str]) -> Optional[OCRResult]:
    print("Engine: Attempting Typhoon OCR...")
    # Simulate failure to allow cascade to continue
    return None

def try_gpt_vision(image_paths: list[str]) -> Optional[OCRResult]:
    print("Engine: Attempting GPT Vision...")
    # Simulate failure
    return None

def try_azure_docint(image_paths: list[str]) -> Optional[OCRResult]:
    print("Engine: Attempting Azure Document Intelligence...")
    # Simulate success for demonstration
    print("Engine: Azure Document Intelligence succeeded.")
    return OCRResult(
        status="OCR_DONE",
        avg_confidence=0.92,
        pages=[{"page_number": 1, "text": "Extracted text from Azure for the image."}],
        tables=[],
        raw_engine_trace={"azure_docint": "success"}
    )

def try_tesseract(image_paths: list[str]) -> Optional[OCRResult]:
    print("Engine: Attempting Tesseract...")
    return None

def try_easyocr(image_paths: list[str]) -> Optional[OCRResult]:
    print("Engine: Attempting EasyOCR...")
    return None


# --- Core OCR Pipeline Logic ---

def run_image_ocr_cascade(image_paths: list[str]) -> OCRResult:
    """
    Runs the cascading OCR pipeline on a list of image file paths.
    """
    trace = {}

    # 1. Typhoon OCR
    result = try_typhoon_ocr(image_paths)
    trace["typhoon_ocr"] = "attempted"
    if result and result.avg_confidence >= 0.8:
        trace["typhoon_ocr"] = "success"
        result.raw_engine_trace.update(trace)
        return result

    # 2. GPT Vision
    result = try_gpt_vision(image_paths)
    trace["gpt_vision"] = "attempted"
    if result and result.avg_confidence >= 0.75:
        trace["gpt_vision"] = "success"
        result.raw_engine_trace.update(trace)
        return result

    # 3. Azure Document Intelligence
    result = try_azure_docint(image_paths)
    trace["azure_docint"] = "attempted"
    if result and result.avg_confidence >= 0.75:
        trace["azure_docint"] = "success"
        result.raw_engine_trace.update(trace)
        return result

    # ... (Tesseract and EasyOCR would follow here) ...

    trace["final_status"] = "all engines failed"
    return OCRResult(
        status="FAILED_OCR",
        avg_confidence=0.0,
        pages=[],
        tables=[],
        raw_engine_trace=trace
    )


def run_cascading_ocr(file_path: str, file_extension: str) -> OCRResult:
    """
    Main function to orchestrate OCR based on file type.
    """
    ext = file_extension.lower().strip(".")
    N_MIN_CHARS_PER_PAGE = 50 # Threshold to decide if a PDF is native or scanned

    if ext == "pdf":
        try:
            with pdfplumber.open(file_path) as pdf:
                full_text = []
                tables = []
                total_chars = 0
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    total_chars += len(text)
                    full_text.append({"page_number": i + 1, "text": text})

                    extracted_tables = page.extract_tables()
                    for tbl in extracted_tables:
                        tables.append({"page_number": i + 1, "cells": tbl})

                # Heuristic: if average characters per page is high, it's a native PDF
                if total_chars / len(pdf.pages) > N_MIN_CHARS_PER_PAGE:
                    print(f"File '{os.path.basename(file_path)}' detected as NATIVE PDF.")
                    return OCRResult(
                        status="OCR_DONE",
                        avg_confidence=0.95,
                        pages=full_text,
                        tables=tables,
                        raw_engine_trace={"engine": "pdfplumber", "strategy": "native_text"}
                    )
                else:
                    # Fallback to image OCR
                    print(f"File '{os.path.basename(file_path)}' detected as SCANNED PDF. Cascading...")
                    # This part needs a PDF-to-image conversion library like pdf2image
                    # For now, we'll simulate this and call the image cascade.
                    # images = convert_pdf_to_images(file_path)
                    return run_image_ocr_cascade([file_path]) # Simulate with path
        except Exception as e:
            return OCRResult(status="FAILED_OCR", avg_confidence=0, pages=[], tables=[], raw_engine_trace={"pdfplumber_error": str(e)})

    elif ext == "docx":
        try:
            document = docx.Document(file_path)
            full_text = "\n".join([p.text for p in document.paragraphs])
            pages = [{"page_number": 1, "text": full_text}]
            tables = []
            for i, table in enumerate(document.tables):
                cells = [[cell.text for cell in row.cells] for row in table.rows]
                tables.append({"page_number": 1, "cells": cells})

            return OCRResult(
                status="OCR_DONE",
                avg_confidence=0.98,
                pages=pages,
                tables=tables,
                raw_engine_trace={"engine": "python-docx"}
            )
        except Exception as e:
            return OCRResult(status="FAILED_OCR", avg_confidence=0, pages=[], tables=[], raw_engine_trace={"docx_error": str(e)})

    elif ext in {"png", "jpg", "jpeg", "tiff", "bmp", "webp"}:
        return run_image_ocr_cascade([file_path])

    else:
        return OCRResult(
            status="FAILED_OCR",
            avg_confidence=0.0,
            pages=[],
            tables=[],
            raw_engine_trace={"error": f"Unsupported file extension: {ext}"}
        )
